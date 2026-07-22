import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('AI_Trainer_BaoCao_TongHop.docx')

def set_style_id(paragraph, style_id):
    pPr = paragraph._p.get_or_add_pPr()
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        pPr.insert(0, pStyle)
    pStyle.set(qn('w:val'), style_id)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text)
    style_map = {1: 'Heading1', 2: 'Heading2', 3: 'Heading3'}
    set_style_id(p, style_map.get(level, 'Heading1'))
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def add_bold_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Paragraph')
    return p

def add_table_with_borders(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = 'Table Grid'
    except:
        tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)
        tblPr.append(borders)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table

# ============================================================
# PHẦN 5: TÀI LIỆU THAM KHẢO – MINDFLOW CHI TIẾT
# ============================================================
doc.add_page_break()
add_heading(doc, 'PHẦN 5: TÀI LIỆU THAM KHẢO – CÁC BLOCK TRONG MINDFLOW (CHI TIẾT)', 1)
add_para(doc, 'Nguồn: Tài liệu "Conductify Mindflow - AI Agent v1.0" – Conductify AI.')
add_para(doc, 'Phần này cung cấp hướng dẫn chi tiết về từng block Mindflow thường được sử dụng nhất để thiết kế AI Agent, bao gồm cách thiết lập và các lưu ý quan trọng.')

# --- 5.1 Session Manager ---
add_heading(doc, '5.1. Session Manager', 2)
add_para(doc, 'Session Manager mặc định là block đầu tiên trong Mindflow. Block này giúp AI quản lý trạng thái sau mỗi lần hỏi đáp. Thông thường, block này sẽ được để trống, không chỉnh sửa gì thêm.')

add_heading(doc, '5.1.1. Extra Fields', 3)
add_para(doc, 'Khi cần thiết, có thể thêm các biến theo định dạng JSON vào khung Extra Fields. Khi cần sử dụng biến đó ở một block khác, gọi biến theo cú pháp:')
add_bold_para(doc, '{{ session.extra.ten_bien }}')
add_para(doc, 'Ví dụ: Thêm biến {"customer_type": ""} vào Extra Fields, sau đó gọi {{ session.extra.customer_type }} ở các block phía sau để lưu trữ và truy xuất loại khách hàng.')

add_heading(doc, '5.1.2. Kiểm tra với Raw Output', 3)
add_para(doc, 'Để kiểm tra, chọn Show Raw Output và nhắn một câu bên section Testing Playground để chạy vào đúng luồng vừa thiết kế. Bên dưới "Raw Output from Mindflow Executor" sẽ hiển thị đoạn prompt từ Mindflow.')

# --- 5.2 Router ---
add_heading(doc, '5.2. Router', 2)
add_para(doc, 'Router được sử dụng để định tuyến luồng. Khi tạo một Agent mới, mặc định Agent sẽ có 4 luồng suy nghĩ (greetings, qna, off_topic và handover). Chúng được gọi là Flow. Mỗi Flow chứa các logic để xử lý riêng một chủ đề công việc gọi là Topic.')

add_heading(doc, '5.2.1. Cách thiết lập', 3)
add_bullet(doc, 'Chọn Add topic để thêm một chủ đề mới.')
add_bullet(doc, 'Chọn Delete topic để xóa một topic không sử dụng.')

add_heading(doc, '5.2.2. Quy tắc đặt tên Topic', 3)
add_bullet(doc, 'Viết liền không dấu (ví dụ: greetings, off_topic, company_info,...)')
add_bullet(doc, 'Có thể thêm bao nhiêu topic tùy mục đích sử dụng.')
add_bullet(doc, 'Nên đặt tên topic có liên quan đến chủ đề thực tế. Tránh đặt những tên không liên quan như sandbox, test, everything,...')

add_heading(doc, '5.2.3. Phrases và Keywords', 3)
add_para(doc, 'Mỗi topic sau khi tạo sẽ có 2 mục Phrases và Keywords. Đây là thông tin đầu vào để giúp AI Tổng quản phân tích ngữ nghĩa câu nói của User và phân loại vào chủ đề thích hợp.')
add_bullet(doc, 'Phrases: chứa những câu hỏi, câu nói tự nhiên của User.')
add_bullet(doc, 'Keywords: chứa những từ khóa để giúp AI nhận biết.')
add_bold_para(doc, '→ Lưu ý: Hạn chế sử dụng keyword. Ưu tiên dùng Phrases.')

# --- 5.3 Flows Switch ---
add_heading(doc, '5.3. Flows Switch', 2)
add_para(doc, 'Flows Switch là block tiếp nối với Router, dùng để tạo Flow và gán Topic nào thuộc Flow nào.')

add_heading(doc, '5.3.1. Cách thiết lập', 3)
add_bullet(doc, 'Chọn Add Flow để thêm một luồng mới.')
add_bullet(doc, 'Chọn Delete Flow để xóa một flow không sử dụng.')
add_bullet(doc, 'Đặt tên flow: Viết liền không dấu. Nên đặt tên flow giống như tên topic để tránh nhầm lẫn.')

add_heading(doc, '5.3.2. Gán Topic vào Flow', 3)
add_bullet(doc, 'Nhập topic đã tạo từ Router vào mục Topics.')
add_bullet(doc, 'Có thể gán nhiều topic vào một flow. Tuy nhiên, để có kết quả tốt nhất, chỉ nên gán 1 topic cho 1 flow.')
add_bullet(doc, 'Chọn biểu tượng mũi tên ở ô "Available Topics can use" để chọn topic từ Router.')

add_bold_para(doc, 'Quy tắc TUYỆT ĐỐI:')
add_bullet(doc, 'TUYỆT ĐỐI không nhập một topic không có ở Router.')
add_bullet(doc, 'TUYỆT ĐỐI không được để trống Topics.')

# --- 5.4 AI QnA ---
add_heading(doc, '5.4. AI QnA (Chi tiết)', 2)
add_para(doc, 'AI QnA là block có tích hợp sẵn các model AI. Đây được xem như từng AI Chuyên trách để tập trung xử lý một công việc cụ thể.')

add_heading(doc, '5.4.1. Các trường thiết lập chính', 3)
add_table_with_borders(doc,
    ['Trường', 'Mô tả', 'Lưu ý'],
    [
        ['Name', 'Đặt tên cho block', 'Nên đặt tên mô tả chức năng'],
        ['Instruction Prompt', 'Prompt hướng dẫn cho AI', 'Quan trọng nhất – quyết định hành vi AI'],
        ['Enable Knowledge Aware', 'Bật/tắt sử dụng Knowledge Bank', 'Khi bật, bắt buộc điền Dataset UUID và QnA UUID'],
        ['Knowledge Dataset UUID', 'UUID của bộ Knowledge Book', 'Lấy từ Knowledge Bank > General Knowledge'],
        ['Knowledge Dataset QA UUID', 'UUID của bộ QnA Library', 'Lấy từ Knowledge Bank > Questions & Answers'],
        ['Knowledge Result Limit', 'Giới hạn số kết quả truy xuất', 'VD: giá trị 2 → lấy 2 cặp QnA liên quan nhất'],
        ['Enable Chat History Aware', 'Nhận biết lịch sử chat', 'Giúp AI hiểu ngữ cảnh cuộc hội thoại'],
        ['Enable Human Input Aware', 'Nhận biết câu chat đầu vào', 'Cho phép AI đọc tin nhắn mới nhất của user'],
        ['Variable to save AI reply', 'Đặt tên biến cho phản hồi', 'Gọi lại bằng {{ ten_bien }}'],
    ]
)

add_heading(doc, '5.4.2. Temperature', 3)
add_para(doc, 'Temperature là tham số điều chỉnh độ chính xác/sáng tạo của AI. Giá trị từ 0.0 đến 2.0.')
add_table_with_borders(doc,
    ['Giá trị', 'Đặc điểm', 'Khuyến nghị'],
    [
        ['0.6 trở xuống', 'AI ít sáng tạo, bám sát prompt/dữ liệu', 'Dùng cho tasks cần chính xác cao'],
        ['0.7', 'Cân bằng giữa chính xác và sáng tạo', 'Giá trị đề xuất cho hầu hết trường hợp'],
        ['0.8 trở lên', 'AI phản hồi sáng tạo hơn', 'Cẩn thận với thông tin có thể bịa'],
        ['1.0 trở lên', 'KHÔNG khuyến khích', 'AI có thể trả lời thiếu kiểm soát'],
    ]
)

add_heading(doc, '5.4.3. Max Tokens', 3)
add_para(doc, 'Lượng token tối đa cho phản hồi của AI. Mặc định: 1000 token.')
add_bullet(doc, 'Không nên để giá trị quá cao → ảnh hưởng chi phí sử dụng AI.')
add_bullet(doc, 'Không nên để giá trị quá thấp → AI trả lời chưa hết câu đã vượt lượng token.')
add_bullet(doc, 'Tùy vào mục đích sử dụng để điều chỉnh phù hợp.')

add_heading(doc, '5.4.4. Các model AI hỗ trợ', 3)
add_table_with_borders(doc,
    ['Model', 'Ghi chú'],
    [
        ['GPT-4', 'Nâng cấp tài khoản để sử dụng'],
        ['GPT-4o-mini', 'Model mặc định phổ biến'],
        ['GPT-3.5', 'Model cơ bản'],
        ['Gemini-pro', 'Model từ Google'],
        ['Claude-3-haiku', 'Model từ Anthropic'],
    ]
)

add_heading(doc, '5.4.5. JSON Mode', 3)
add_para(doc, 'Enable JSON Mode: Bật/tắt tính năng buộc AI trả lời theo định dạng JSON. Tính năng này chỉ sử dụng được với các model GPT của OpenAI (GPT-4, GPT-4o-mini, GPT-3.5).')

# --- 5.5 Bot Send Text ---
add_heading(doc, '5.5. Bot Send Text', 2)
add_para(doc, 'Bot Send Text được dùng để thiết lập sẵn câu trả lời mong muốn. Có thể kết hợp sử dụng cú pháp {{ ten_bien }} để gọi biến từ các block trước đó.')

# --- 5.6 Filter Bot Reply ---
add_heading(doc, '5.6. Filter Bot Reply (Chi tiết)', 2)
add_para(doc, 'Filter Bot Reply dùng để thay đổi văn bản (text) do AI trả lời ra.')
add_para(doc, 'Ví dụ: AI xưng "Tôi" khi trả lời với User, nhưng mong muốn phải là xưng "Em".')
add_bold_para(doc, 'Điều kiện sử dụng: Phải sử dụng block AI QnA trước đó.')

add_heading(doc, '5.6.1. Cách thiết lập', 3)
add_bullet(doc, 'Regex: Nhập văn bản cần thay thế theo cú pháp (?:text). Nếu có nhiều hơn một text: (?:text1|text2|text3).')
add_bullet(doc, 'Replace: Nhập văn bản mong muốn thay thế.')
add_para(doc, 'Ví dụ: Regex = (?:Tôi|tôi) → Replace = Em. Kết quả: mọi chỗ AI xưng "Tôi" hoặc "tôi" sẽ tự động đổi thành "Em".')

# --- 5.7 Clear Topic ---
add_heading(doc, '5.7. Clear Topic', 2)
add_para(doc, 'Clear topic thường được đặt ở cuối mỗi flow, dùng để thiết lập lại flow, xóa topic để câu nói tiếp theo của User quay lại Router và tiếp tục được phân luồng từ đầu.')
add_para(doc, 'Block này chỉ cần kéo vào sử dụng, không cần thiết lập gì thêm.')

# --- 5.8 Break ---
add_heading(doc, '5.8. Break', 2)
add_para(doc, 'Break được dùng để ngừng tất cả các block bên dưới. Ví dụ: nếu đặt Break trước Filter Bot Reply và Clear topic thì 2 block đó sẽ không được thực thi.')
add_bold_para(doc, '→ Break thường được dùng kết hợp với Bot Send Text trong việc debug.')

# --- 5.9 Tổng kết ---
add_heading(doc, '5.9. Tổng kết: Thứ tự block phổ biến trong một Flow', 2)
add_para(doc, 'Dưới đây là thứ tự các block thường gặp nhất trong một flow hoàn chỉnh:')
add_table_with_borders(doc,
    ['Thứ tự', 'Block', 'Chức năng'],
    [
        ['1', 'Session Manager', 'Quản lý trạng thái, lưu biến session'],
        ['2', 'Router', 'Phân tích ngữ nghĩa, phân loại topic'],
        ['3', 'Flows Switch', 'Chuyển hướng đến flow tương ứng'],
        ['4', 'AI QnA', 'AI Chuyên trách xử lý, trả lời'],
        ['5', 'Filter Bot Reply', 'Chỉnh sửa văn bản phản hồi (nếu cần)'],
        ['6', 'Bot Send Text', 'Gửi câu trả lời cho User'],
        ['7', 'Clear Topic', 'Reset flow, quay lại Router cho câu tiếp'],
    ]
)

# ============================================================
# PHẦN 6: TÀI LIỆU THAM KHẢO – KNOWLEDGE BANK
# ============================================================
doc.add_page_break()
add_heading(doc, 'PHẦN 6: TÀI LIỆU THAM KHẢO – CONDUCTIFY STUDIO KNOWLEDGE BANK', 1)
add_para(doc, 'Nguồn: Tài liệu "Conductify Studio - Knowledge Bank v1.0" – Conductify AI.')
add_para(doc, 'Phần này hướng dẫn chi tiết cách sử dụng Knowledge Bank – nơi lưu trữ và quản lý kiến thức cho AI Agent trên nền tảng Conductify Studio.')

# --- 6.1 Tổng quan ---
add_heading(doc, '6.1. Tổng quan về Knowledge Bank', 2)
add_para(doc, 'Bản thân mỗi model AI đã được huấn luyện sẵn một lượng kiến thức nền đủ lớn. Tuy nhiên, khi ứng dụng vào môi trường doanh nghiệp, AI cần được bổ sung kiến thức chuyên môn riêng biệt. Knowledge Bank chính là nơi lưu trữ và quản lý những kiến thức này.')

add_heading(doc, '6.1.1. Hai loại kiến thức trong hệ thống', 3)
add_table_with_borders(doc,
    ['Loại kiến thức', 'Dạng dữ liệu', 'Khi nào sử dụng', 'Độ ưu tiên'],
    [
        ['General Knowledge', 'Chủ đề – Nội dung (Book)', 'Nội dung thông tin lớn, cố định', 'Thấp hơn QnA'],
        ['Questions & Answers (QnA)', 'Câu hỏi – Câu trả lời', 'Nội dung thông tin nhỏ, cụ thể', 'Cao hơn (AI ưu tiên trả lời sát QnA)'],
    ]
)
add_bold_para(doc, 'Lưu ý quan trọng: AI sẽ ưu tiên trả lời sát theo nội dung trong QnA hơn so với nội dung trong General Knowledge.')

# --- 6.2 General Knowledge ---
add_heading(doc, '6.2. General Knowledge', 2)

add_heading(doc, '6.2.1. Cách tạo một Book mới', 3)
add_para(doc, 'Tại giao diện Knowledge Bank → tab General Knowledge → chọn Create new Book.')
add_bold_para(doc, 'Các trường thiết lập:')
add_table_with_borders(doc,
    ['Trường', 'Mô tả', 'Lưu ý'],
    [
        ['Book Name', 'Đặt tên cho Book', 'Viết liền không dấu, tối đa 30 ký tự (VD: vhn_product_info)'],
        ['Access', 'Phạm vi chia sẻ', 'Only me (mặc định) hoặc All team members'],
        ['AI Content Analyser', 'Model AI phân tích văn bản', 'OpenAI (hiểu tiếng Anh tốt) hoặc VnEm (hiểu tiếng Việt tốt). KHÔNG dùng 8kEm'],
    ]
)

add_heading(doc, '6.2.2. Cách thêm dữ liệu vào Book (Datasource)', 3)
add_para(doc, 'Chọn Book → tab Datasources → Add new datasource.')
add_bold_para(doc, 'Các Source format hỗ trợ:')
add_table_with_borders(doc,
    ['Source format', 'Mô tả', 'Lưu ý'],
    [
        ['File', 'Kéo thả hoặc chọn file (pdf, docx,...)', 'Chỉ trích xuất văn bản, không lấy được hình ảnh/video'],
        ['Text', 'Nhập nội dung dữ liệu trực tiếp', 'Khuyến khích kiểm tra và bổ sung sau khi dùng File/URL'],
        ['URL', 'Nhập đường link trang web', 'Có thể nhập nhiều URL, mỗi URL trên 1 dòng'],
        ['Crawled', 'Không sử dụng', '—'],
    ]
)
add_bold_para(doc, 'Lưu ý quan trọng:')
add_bullet(doc, 'Đối với file và url, sau khi xử lý xong, nội dung sẽ hiển thị ở phần Text.')
add_bullet(doc, 'Hệ thống chỉ hỗ trợ trích xuất văn bản. Hình ảnh, video sẽ không lấy được.')
add_bullet(doc, 'Nên vào phần Text để chỉnh sửa, bổ sung mô tả cho các nội dung hình ảnh/bảng biểu.')
add_bullet(doc, 'Chọn Save & Process để lưu. Chọn Reload list để xem kết quả sau khi hệ thống xử lý xong.')

# --- 6.3 Questions & Answers ---
add_heading(doc, '6.3. Questions & Answers (QnA)', 2)

add_heading(doc, '6.3.1. Cách tạo Q&A Library mới', 3)
add_para(doc, 'Tại giao diện Knowledge Bank → tab Questions & Answers → chọn Create new Q&A Library.')
add_table_with_borders(doc,
    ['Trường', 'Mô tả', 'Lưu ý'],
    [
        ['Q&A Library Name', 'Đặt tên cho Library', 'Viết liền không dấu, tối đa 30 ký tự'],
        ['Access', 'Phạm vi chia sẻ', 'Only me (mặc định) hoặc All team members'],
        ['Language Model', 'Model AI phân tích', 'Cố định là VnEm (không thay đổi được)'],
    ]
)

add_heading(doc, '6.3.2. Cách thêm dữ liệu vào Q&A Library', 3)
add_para(doc, 'Chọn Q&A Library → tab Q&A Pairs → Add new Q&A Pairs.')
add_bullet(doc, 'Nhập câu hỏi vào khung Question.')
add_bullet(doc, 'Nhập câu trả lời vào khung Answer.')
add_bullet(doc, 'Chọn Remember → Confirm để lưu.')

add_heading(doc, '6.3.3. Import từ CSV (Quan trọng)', 3)
add_bold_para(doc, 'Đối với bộ dữ liệu lớn, nhiều cặp Q&A, PHẢI dùng chức năng Import from CSV để tiết kiệm thời gian.')
add_para(doc, 'Các bước thực hiện:')
add_bullet(doc, 'Bước 1: Tạo bảng tính với 2 cột: question và answer. Nhập dữ liệu từng cặp theo từng dòng.')
add_bullet(doc, 'Bước 2: Tải bảng tính về với định dạng CSV.')
add_bullet(doc, 'Bước 3: Trong Q&A Library, chọn Import from CSV.')
add_bullet(doc, 'Bước 4: Kéo thả hoặc chọn file CSV → Import → Confirm.')
add_bullet(doc, 'Bước 5: Chọn Reload list để hiển thị kết quả.')
add_para(doc, 'Ngược lại, khi cần trích dữ liệu từ hệ thống thành file CSV, chọn Export to CSV.')

# --- 6.4 Quick Test ---
add_heading(doc, '6.4. Kiểm tra nhanh (Quick Test)', 2)
add_para(doc, 'Trong mỗi Book hoặc Q&A Library, sau khi nạp dữ liệu, có thể kiểm tra nhanh khả năng truy vấn dữ liệu bằng tab Quick Test.')
add_bold_para(doc, 'Cách sử dụng:')
add_bullet(doc, 'Nhập câu hỏi ở ô Your Question.')
add_bullet(doc, 'Nhập giới hạn số lượng kết quả tại ô Limit number of matches to retrieve.')
add_bullet(doc, 'Chọn Test question và xem kết quả bên phía tay phải.')

# --- 6.5 Content Splitter ---
add_heading(doc, '6.5. Cách thức hoạt động của Datasource – Content Splitter', 2)
add_para(doc, 'Dù nhập dữ liệu bằng file, text hay url, hệ thống sẽ cắt nội dung thành từng phần (segment) nhỏ. Ba cài đặt quyết định cách cắt:')

add_heading(doc, '6.5.1. Content Splitter – Cách thức chia nội dung', 3)
add_table_with_borders(doc,
    ['Phương thức', 'Mô tả', 'Chi tiết'],
    [
        ['Markdown', 'Cắt theo markdown', 'Cắt theo tiêu đề (#), khối mã (```), đường kẻ (---), dòng trống, xuống dòng, khoảng trắng'],
        ['Paragraph', 'Cắt theo đoạn văn', 'Tách nội dung theo từng đoạn văn bản'],
        ['Token', 'Cắt theo số lượng token', 'Token = các từ/cụm từ được tách ra để phân tích'],
    ]
)

add_heading(doc, '6.5.2. Maximum Segment Length và Overlap', 3)
add_para(doc, 'Hai trường này có tác dụng khi sử dụng Content Splitter bằng Token:')
add_bullet(doc, 'Maximum Segment Length: Độ dài tối đa cho 1 segment. VD: text gồm 1450 token, segment length = 700 → tách thành 3 segment (700 + 700 + 50).')
add_bullet(doc, 'Overlap length between segments: Độ dài chồng chéo giữa các segment. VD: text gồm 1200 token, segment = 700, overlap = 200 → segment 1 gồm 700 token đầu, segment 2 gồm 700 token (bắt đầu từ token thứ 501).')
add_para(doc, 'Mục đích của overlap: đảm bảo thông tin ở ranh giới giữa 2 segment không bị mất, giúp AI truy xuất ngữ cảnh đầy đủ hơn.')

# --- 6.6 Kết nối Knowledge Bank với AI QnA ---
add_heading(doc, '6.6. Kết nối Knowledge Bank với Mindflow (AI QnA)', 2)
add_para(doc, 'Để AI Agent sử dụng được kiến thức từ Knowledge Bank, cần thiết lập trong block AI QnA:')
add_bullet(doc, 'Bật Enable Knowledge Aware.')
add_bullet(doc, 'Nhập Knowledge Dataset UUID: lấy từ tab Knowledge Book trong General Knowledge.')
add_bullet(doc, 'Nhập Knowledge Dataset QA UUID: lấy từ tab Q&A Settings trong Q&A Library.')
add_bullet(doc, 'Thiết lập Knowledge Result Limit phù hợp (khuyến nghị: 3-5 cho Book, 2-3 cho QnA).')
add_bold_para(doc, 'Lưu ý: Nếu bật Knowledge Aware mà không điền đủ UUID, AI sẽ không thể truy xuất kiến thức và có thể trả lời sai.')

# --- 6.7 Tổng kết ---
add_heading(doc, '6.7. Tổng kết Knowledge Bank', 2)
add_table_with_borders(doc,
    ['Hạng mục', 'General Knowledge (Book)', 'Questions & Answers (QnA)'],
    [
        ['Dạng dữ liệu', 'Chủ đề – Nội dung dài', 'Câu hỏi – Câu trả lời ngắn'],
        ['Nguồn nhập', 'File, Text, URL', 'Thủ công hoặc Import CSV'],
        ['AI Analyser', 'OpenAI hoặc VnEm', 'VnEm (cố định)'],
        ['Chia nhỏ dữ liệu', 'Markdown / Paragraph / Token', 'Không áp dụng (mỗi cặp Q&A là 1 đơn vị)'],
        ['Độ ưu tiên AI', 'Thấp hơn', 'Cao hơn (AI ưu tiên QnA)'],
        ['UUID cần lấy', 'Dataset UUID', 'Dataset QnA UUID'],
        ['Quick Test', 'Có', 'Có'],
    ]
)

# Save
doc.save('AI_Trainer_BaoCao_TongHop.docx')
print('OK - Saved successfully')
print(f'Total paragraphs: {len(doc.paragraphs)}')
