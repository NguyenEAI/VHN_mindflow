import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('AI_Trainer_BaoCao_TongHop.docx')

def set_style_id(paragraph, style_id):
    pPr = paragraph._p.get_or_add_pPr()
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        pPr.insert(0, pStyle)
    pStyle.set(qn('w:val'), style_id)

# Override add_heading to use style_id
def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_style_id(p, f'Heading{level}')
    return p

# Helpers
def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_bullet(text):
    p = doc.add_paragraph()
    p.add_run(text)
    set_style_id(p, 'ListParagraph')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    try:
        table.style = 'Table Grid'
    except KeyError:
        # Set table borders manually via XML
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = _OE('w:tblBorders')
        for b in ('top','left','bottom','right','insideH','insideV'):
            el = _OE(f'w:{b}')
            el.set(_qn('w:val'),'single')
            el.set(_qn('w:sz'),'4')
            el.set(_qn('w:color'),'808080')
            tblBorders.append(el)
        tblPr.append(tblBorders)
    hdr = table.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs:
            for run in r.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    return table

# ===== PHẦN 4: BUỔI 11 =====
doc.add_page_break()
add_heading('PHẦN 4: BUỔI 11 – AI SALES AGENT & TÍCH HỢP ĐA KÊNH', level=1)
add_para('Advance Level')

# 4.1
add_heading('4.1. Use Case: AI Sales Agent', level=2)

add_para('Mục đích:', bold=True)
add_para('Hỗ trợ doanh nghiệp bán hàng, tư vấn sản phẩm theo nhu cầu của khách hàng, gợi ý sản phẩm phù hợp, chốt đơn và ghi nhận thông tin đơn hàng vào hệ thống nội bộ.')

add_para('Chức năng chính:', bold=True)
add_bullet('Thông tin chung: Cung cấp thông tin về doanh nghiệp, chi nhánh, chính sách bán hàng, chương trình khuyến mãi hiện hành.')
add_bullet('Tư vấn sản phẩm: Lắng nghe nhu cầu khách hàng để gợi ý sản phẩm phù hợp, kết hợp upsale và cross-sale.')
add_bullet('Hỗ trợ đặt hàng: Thêm/bớt sản phẩm vào giỏ, tính toán tổng đơn hàng theo chính sách doanh nghiệp, chốt đơn và ghi nhận vào hệ thống quản trị nội bộ.')

add_para('Kỳ vọng:', bold=True)
add_bullet('Agent đối đáp và tư vấn tự nhiên, linh hoạt đề xuất sản phẩm phù hợp.')
add_bullet('Agent chốt được đơn và ghi nhận đơn hàng vào hệ thống quản trị nội bộ.')
add_bullet('Agent phản hồi với ngôn ngữ lịch sự như một nhân viên sales thực thụ.')
add_bullet('Agent tạo ticket để gọi nội bộ hỗ trợ khi không xử lý được yêu cầu.')

add_para('Ứng dụng thực tế:', bold=True)
add_para('AI Sales Agent của một công ty Affiliate, bán hàng cho các thương hiệu mỹ phẩm trong nước (Shegan, Cocoon, Ofelia). Hiện đang vận hành 24/7 trên 8 tài khoản Zalo cá nhân, 5 Fanpage Facebook và 6 Website/Landing page.')

# 4.2
add_heading('4.2. Quy trình thiết kế AI Sales Agent', level=2)

add_heading('4.2.1. Flowchart tổng thể', level=3)
add_para('AI Sales Agent gồm các luồng chính:')
add_bullet('Các luồng hỏi đáp thông tin: Chào hỏi, Off-topic, Thông tin chung, Chi nhánh, Khuyến mãi.')
add_bullet('Luồng Gọi hỗ trợ (Handover).')
add_bullet('Luồng Tư vấn sản phẩm.')
add_bullet('Luồng Đặt hàng.')

add_heading('4.2.2. Mindflow – Các luồng hỏi đáp thông tin', level=3)
add_para('Với nhóm luồng Chào hỏi, Off-topic, Thông tin chung, Chi nhánh, Khuyến mãi: sử dụng các kỹ thuật prompting và xử lý dữ liệu tại block AI QnA. Đây là phần Basic, tận dụng Knowledge Dataset (General Knowledge + Q&A) đã chuẩn bị.')

add_heading('4.2.3. Mindflow – Luồng Gọi hỗ trợ (Handover)', level=3)
add_para('Sử dụng các kỹ thuật nâng cao:')
add_bullet('AI Phân tích để trả về kết quả dạng JSON.')
add_bullet('Thao tác với biến: system variables kết hợp json.loads để parse JSON.')
add_bullet('Gọi API vào hệ thống nội bộ để tạo ticket hỗ trợ.')

add_heading('4.2.4. Mindflow – Luồng Tư vấn sản phẩm', level=3)
add_para('Logic: AI Agent vấn đáp với khách hàng để tìm hiểu nhu cầu, phân tích thông tin thành các biến cần thiết để truy vấn Database sản phẩm (AI Vector DB), từ đó có thêm thông tin để tư vấn chính xác hơn.')
add_para('Các kỹ thuật chính:', bold=True)
add_table(['STT', 'Kỹ thuật'], [
    ['1', 'AI Phân tích: trích xuất nhu cầu thành các biến cần thiết.'],
    ['2', 'Block Python: kiểm tra đã đủ các biến cần thiết chưa.'],
    ['3', 'If + Recall: nếu chưa đủ thì chạy lại để hỏi tiếp.'],
    ['4', 'Bot API Call: gọi API vào AI Vector DB.'],
    ['5', 'Sử dụng kết quả từ API để trả lời khách hàng.'],
    ['6', 'Lưu lịch sử chat (Session Extra).'],
    ['7', 'Lưu tên sản phẩm để tái sử dụng ở luồng Đặt hàng.'],
])

add_heading('4.2.5. Mindflow – Luồng Đặt hàng', level=3)
add_para('Logic: AI Agent dùng thông tin sản phẩm đã lưu ở luồng Tư vấn để gọi lại Database lấy thêm thông tin báo giá và đề xuất combo liên quan. Khi khách chốt đơn, AI Agent hỏi thêm thông tin cá nhân (địa chỉ để tính phí ship) và ghi nhận vào hệ thống nội bộ.')

add_para('A. AI Tổng quản', bold=True)
add_bullet('AI Tổng quản theo dõi cuộc trò chuyện, phân tích thông tin thành JSON. Lưu ý: KHÔNG bật Chat History Aware và Human Input Aware để tránh ghi nhớ đơn hàng cũ.')
add_bullet('Block Python: kiểm tra đã đủ các biến cần thiết chưa.')
add_bullet('Block If: nếu chưa đủ biến → Recall để phân tích lại.')

add_para('B. Gọi API AI Vector DB', bold=True)
add_bullet('Block If kiểm tra thông tin sản phẩm đã được cung cấp chưa.')
add_bullet('Nếu khách hàng chưa muốn mua: hỏi thông tin cá nhân để ghi nhận khách hàng tiềm năng. Nếu khách thoát thì thoát luồng và xóa lịch sử chat đơn cũ.')
add_bullet('Nếu đã có thông tin sản phẩm: gọi API Vector DB. Use case Shegan có bán cả sản phẩm lẻ và combo nên gọi 2 lần để báo giá sản phẩm + gợi ý combo liên quan.')

add_para('C. Xử lý kết quả search', bold=True)
add_bullet('Block If kiểm tra gọi API có thành công không.')
add_bullet('Nếu lỗi: ghi nhận lỗi và thử lại.')
add_bullet('Nếu thành công: dùng Python để làm gọn kết quả search trước khi đưa cho AI Chuyên trách.')

add_para('D. Tính toán giá trị đơn hàng', bold=True)
add_bullet('AI Phân tích kiểm tra khách hàng đã cung cấp thông tin địa chỉ chưa.')
add_bullet('Block Python check biến + Block If phân luồng.')
add_bullet('Nếu chưa đủ → Recall để hỏi lại.')
add_bullet('Nếu đủ → áp dụng chính sách bán hàng để tính giá tổng đơn hàng.')

add_para('E. Ghi nhận đơn hàng', bold=True)
add_bullet('Thu thập thông tin cá nhân và ý định khách hàng (chốt đơn / thay đổi / dừng mua) – giống luồng Register.')
add_bullet('Logic phân luồng giống bước 2 của phần B.')
add_bullet('Gọi API ghi nhận đơn hàng vào hệ thống nội bộ.')
add_bullet('Trả lời xác nhận cho khách hàng.')
add_bullet('Xóa lịch sử chat đơn cũ, xóa thông tin sản phẩm để chuẩn bị cho đơn tiếp theo.')
add_bullet('Lưu thông tin khách hàng để chăm sóc về sau.')

add_para('Lưu ý quan trọng:', bold=True)
add_para('Đây là Mindflow mẫu cho dự án Shegan. Với từng dự án cụ thể khác, cần tùy chỉnh logic và kỹ thuật cho phù hợp với đặc thù sản phẩm/chính sách doanh nghiệp.')

# 4.3
add_heading('4.3. Integration Browser – Tích hợp AI Agent đa kênh', level=2)

add_heading('4.3.1. Chức năng', level=3)
add_para('Integration Browser là giải pháp tích hợp AI Agent đã thiết kế lên các nền tảng chat như Zalo và Facebook Messenger, cho phép vận hành đa kênh từ một AI Agent duy nhất.')

add_heading('4.3.2. Tạo Profiles', level=3)
add_para('Sau khi tải về và mở Integration Browser, trang Profiles xuất hiện. Chọn Add để tạo Profile mới, nhập tên và chọn Create. Có thể chỉnh sửa tên, thêm/xóa profile sau khi tạo. Chọn một profile đã tạo để tiếp tục.')

add_heading('4.3.3. Tích hợp AI Agent lên Zalo', level=3)
add_table(['Bước', 'Thao tác'], [
    ['B1', 'Truy cập https://chat.zalo.me/ và đăng nhập tài khoản Zalo cá nhân.'],
    ['B2', 'Đăng nhập Integration Browser bằng tài khoản Conductify Studio.'],
    ['B3', 'Cấu hình: Active AI Integration + Active Zalo Integration; chọn AI Bot (đã Published); AI Mode = AI Take Control; thiết lập Working Time; nhập số điện thoại của tài khoản Zalo vào Zalo Current Chat Phone Number; cuộn xuống và Save.'],
])

add_heading('4.3.4. Tích hợp AI Agent lên Facebook Messenger', level=3)
add_table(['Bước', 'Thao tác'], [
    ['B1', 'Truy cập https://www.facebook.com/ và đăng nhập tài khoản Facebook cá nhân.'],
    ['B2', 'Truy cập trang Messenger Business theo hướng dẫn.'],
    ['B3', 'Đăng nhập Conductify Studio và cấu hình: Active AI Integration + Active Facebook Integration; chọn AI Bot (đã Published); AI Mode = AI Take Control; thiết lập Working Time; nhập asset_id (hiển thị trên URL) vào Facebook Page ID; nhấn Save.'],
])

# 4.4
add_heading('4.4. AI CRM – Quản lý hội thoại đa kênh', level=2)

add_heading('4.4.1. Chức năng', level=3)
add_para('AI CRM quản lý toàn bộ cuộc trò chuyện của AI Agent trên các nền tảng Zalo, Facebook và Website. Ngoài ra, người dùng có thể chat lại trực tiếp với User thông qua AI CRM (chế độ Manual).')

add_heading('4.4.2. Đăng nhập và Giao diện', level=3)
add_para('Truy cập https://hub.conductify.ai/ – đăng ký/đăng nhập bằng tài khoản đã đăng ký.')
add_para('Dashboard:', bold=True)
add_bullet('Hiển thị các AI Agent đã được thiết kế và Published.')
add_bullet('Biểu tượng bánh răng: vào giao diện quản lý Agent.')
add_bullet('Nút OPEN: mở chat standalone.')
add_bullet('Thanh tìm kiếm AI Agent.')
add_para('Giao diện quản lý Agent gồm 5 tab: CRM, Overview, Q&A, Integrations, Contact list.')

add_heading('4.4.3. Tính năng chính của AI CRM', level=3)

add_para('Tìm kiếm cuộc trò chuyện', bold=True)
add_para('Có sẵn nhiều filter: thanh search, nút filter, chọn khung thời gian, chọn kênh (Zalo / Facebook / Website), chọn Mode.')

add_para('Manual – Người dùng tiếp quản', bold=True)
add_para('Hỗ trợ chế độ Manual cho phép người dùng tự nhắn cho User trong trường hợp AI Agent chưa xử lý được. Có thể bật/tắt chế độ Manual. Tin nhắn được phân loại Send by AI (do AI Agent xử lý) hoặc Send by Human (do người dùng nhắn từ AI CRM). Textbox nhập tay chỉ hiển thị ở chế độ Manual.')

add_para('Overview – Điều chỉnh Agent Profile', bold=True)
add_para('Chỉnh sửa Name, Description, Icon của Agent trực tiếp trên AI CRM, tương đương tab Agent Profile ở Studio.')

add_para('Q&A – Chỉnh sửa kiến thức', bold=True)
add_para('Cho phép điều chỉnh các Q&A trên AI Agent trực tiếp mà không cần vào Studio. Trong từng bộ Q&A có thanh search, thêm mới, chỉnh sửa, xóa cặp Question – Answer.')

add_para('Integrations – Thông tin tích hợp', bold=True)
add_para('Hiển thị link chat standalone (Public URL) và hướng dẫn nhúng AI Agent vào website doanh nghiệp (Web Embed).')

add_para('Contact list – Lưu thông tin đơn hàng', bold=True)
add_para('Nếu chưa có hệ thống quản trị nội bộ, có thể gọi API lưu thông tin đơn hàng trực tiếp trên AI CRM. Các trường mặc định: STT, Ngày tạo, Tên, Email, Số điện thoại, Địa chỉ, Công ty, Thông tin khác. Hỗ trợ thanh search, phân trang, xuất file CSV.')

add_heading('4.4.4. API gọi vào Contact list', level=3)
add_table(['Thành phần', 'Giá trị'], [
    ['URL', 'https://api.workerbot.ai/v1/contact-list/create'],
    ['Method', 'POST'],
    ['Raw Body (JSON)', '{ "email", "name", "phone", "address", "company", "bot_uuid", "thread_id", "full_field": { "product_requirement": "...", "special_requirement": "..." } }'],
])
add_para('Ví dụ Raw Body:', bold=True)
p = doc.add_paragraph()
p.add_run('{\n'
          '  "email": "xyz@gmail.com",\n'
          '  "name": "Nguyễn Văn B",\n'
          '  "phone": "0987654321",\n'
          '  "address": "10/11 Trần Văn Ơn",\n'
          '  "company": "Conductify AI",\n'
          '  "bot_uuid": "871521c6-6b2a-43fa-bb78-7d45c7bce794",\n'
          '  "thread_id": "07156a29-ed14-4883-96cf-e978c8300b1c",\n'
          '  "full_field": { "product_requirement": "2 chai sữa rửa mặt", "special_requirement": "gói quà" }\n'
          '}')
add_para('Lưu ý:', bold=True)
add_bullet('bot_uuid và thread_id là system variables, lấy trực tiếp từ Mindflow.')
add_bullet('Các biến còn lại có thể dùng AI Phân tích để trả về kết quả.')
add_bullet('Trường full_field có giá trị là JSON, có thể thêm/bớt field tùy mục đích sử dụng.')

add_heading('4.4.5. Tạo AI Agent từ AI CRM', level=3)
add_para('Có thể tạo AI Agent trực tiếp từ AI CRM bằng nút Create New Agent ở Dashboard. Có 2 chế độ:')
add_table(['Mode', 'Cách dùng'], [
    ['Easy Mode', 'Điền thông tin sơ bộ cho AI Agent và chờ 72 tiếng để có AI Agent hoàn chỉnh.'],
    ['Expert Mode', 'Điền thông tin và nạp data Q&A để tạo AI Agent từ template sẵn có.'],
])

# 4.5 Summary
add_heading('4.5. Tổng kết Buổi 11', level=2)
add_table(['Hạng mục', 'Nội dung chính'], [
    ['Use case', 'AI Sales Agent: tư vấn, chốt đơn, ghi nhận đơn hàng vào hệ thống.'],
    ['Thiết kế', 'Flowchart + Mindflow với 4 nhóm luồng: hỏi đáp thông tin, gọi hỗ trợ, tư vấn sản phẩm, đặt hàng.'],
    ['Kỹ thuật chính', 'AI Phân tích → JSON, Python kiểm tra biến, If/Recall, Bot API Call (AI Vector DB), lưu Session Extra, Set/Clear Topic.'],
    ['Integration Browser', 'Tích hợp AI Agent lên Zalo cá nhân và Facebook Messenger qua tài khoản Conductify Studio.'],
    ['AI CRM', 'Quản lý hội thoại đa kênh, chế độ Manual, chỉnh sửa Q&A, Contact list, tạo Agent (Easy/Expert mode).'],
    ['API Contact list', 'POST https://api.workerbot.ai/v1/contact-list/create – ghi nhận đơn hàng trực tiếp.'],
])

# Bài thực hành
add_heading('4.6. Bài thực hành', level=2)
add_bullet('Bài 1: Vẽ Flowchart cho AI Sales Agent ứng với doanh nghiệp đang công tác.')
add_bullet('Bài 2: Thiết kế Mindflow chi tiết cho AI Sales Agent này (đầy đủ 4 nhóm luồng).')
add_bullet('Bài 3: Sử dụng Integration Browser để tích hợp AI Agent đã tạo vào Zalo cá nhân và Fanpage doanh nghiệp.')

doc.save('AI_Trainer_BaoCao_TongHop.docx')
print('Saved successfully')
